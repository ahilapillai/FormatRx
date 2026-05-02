"""
exporter.py — Reconstructs a .docx from a (possibly modified) manuscript dict.

Two export modes:
  - clean:    Final document, no tracked-change marks.
  - tracked:  Additions in green underline, deletions in red strikethrough.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import io


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_body(doc: Document, text: str):
    """Add paragraphs from a content block (split on double newlines)."""
    if not text.strip():
        return
    for para_text in text.split("\n\n"):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())


def _set_run_color(run, color: RGBColor):
    run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Tracked change helpers (simple visual approximation via run formatting)
# ---------------------------------------------------------------------------

def _add_tracked_paragraph(doc: Document, original: str, updated: str, inline_diff: list[dict]):
    """
    Add a paragraph showing deletions (red strikethrough) and insertions (green underline).
    Falls back to plain updated text if no inline_diff.
    """
    if not inline_diff:
        doc.add_paragraph(updated)
        return

    para = doc.add_paragraph()
    for segment in inline_diff:
        op = segment["op"]
        text = segment["text"]
        if op == "equal":
            para.add_run(text)
        elif op == "delete":
            run = para.add_run(text)
            run.font.strike = True
            _set_run_color(run, RGBColor(0xCC, 0x00, 0x00))  # Red
        elif op == "insert":
            run = para.add_run(text)
            run.font.underline = True
            _set_run_color(run, RGBColor(0x00, 0x88, 0x00))  # Green


# ---------------------------------------------------------------------------
# Main export
# ---------------------------------------------------------------------------

def export_docx(
    manuscript: dict,
    changes: list[dict],
    mode: str = "clean",  # "clean" | "tracked"
) -> bytes:
    """
    Build and return a .docx file as bytes.

    manuscript: final manuscript dict (with accepted changes already applied for clean mode).
    changes: list of enriched change dicts (with inline_diff) for tracked mode.
    mode: "clean" or "tracked"
    """
    doc = Document()

    # Build a lookup from (section, field) → change for tracked mode
    change_lookup: dict[tuple, dict] = {}
    if mode == "tracked":
        for c in changes:
            key = (c.get("section", "").lower(), c.get("field", ""))
            change_lookup[key] = c

    # ---- Title ----
    title_text = manuscript.get("title", "Untitled")
    doc.add_heading(title_text, 0)

    # ---- Abstract ----
    abstract = manuscript.get("abstract", "")
    if abstract:
        doc.add_heading("Abstract", 1)
        if mode == "tracked":
            key = ("abstract", "content")
            ch = change_lookup.get(key)
            if ch:
                _add_tracked_paragraph(doc, ch["original"], ch["updated"], ch.get("inline_diff", []))
            else:
                doc.add_paragraph(abstract)
        else:
            doc.add_paragraph(abstract)

    # ---- Sections ----
    for sec in manuscript.get("sections", []):
        heading = sec.get("heading", "")
        level = sec.get("level", 1)
        content = sec.get("content", "")

        # Heading
        if mode == "tracked":
            h_key = (heading.lower(), "heading")
            h_ch = change_lookup.get(h_key)
            if h_ch:
                p = doc.add_paragraph()
                p.style = doc.styles[f"Heading {level}"]
                _add_tracked_paragraph(doc, h_ch["original"], h_ch["updated"], h_ch.get("inline_diff", []))
            else:
                doc.add_heading(heading, level)
        else:
            doc.add_heading(heading, level)

        # Content
        if content:
            if mode == "tracked":
                c_key = (heading.lower(), "content")
                c_ch = change_lookup.get(c_key)
                if c_ch:
                    _add_tracked_paragraph(doc, c_ch["original"], c_ch["updated"], c_ch.get("inline_diff", []))
                else:
                    _add_body(doc, content)
            else:
                _add_body(doc, content)

    # ---- References ----
    refs = manuscript.get("references", [])
    if refs:
        doc.add_heading("References", 1)
        for i, ref in enumerate(refs, start=1):
            if mode == "tracked":
                r_key = ("references", f"reference_{i}")
                r_ch = change_lookup.get(r_key)
                if r_ch:
                    _add_tracked_paragraph(doc, r_ch["original"], r_ch["updated"], r_ch.get("inline_diff", []))
                else:
                    doc.add_paragraph(ref)
            else:
                doc.add_paragraph(ref)

    # ---- Figure Legends ----
    figs = manuscript.get("figure_legends", [])
    if figs:
        doc.add_heading("Figure Legends", 1)
        for i, fig in enumerate(figs, start=1):
            if mode == "tracked":
                f_key = ("figure legends", f"figure_{i}")
                f_ch = change_lookup.get(f_key)
                if f_ch:
                    _add_tracked_paragraph(doc, f_ch["original"], f_ch["updated"], f_ch.get("inline_diff", []))
                else:
                    doc.add_paragraph(fig)
            else:
                doc.add_paragraph(fig)

    # ---- Serialize to bytes ----
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
